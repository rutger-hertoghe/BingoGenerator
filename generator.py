import re
import argparse
import random
from docx import Document
from docx.shared import Pt

def SampleInputTxt(inputFile):
  with open(inputFile, 'r', encoding = "utf-8") as file:
    lines = [line.rstrip("\n") for line in file]

  sampleSize = 24
  if len(lines) < sampleSize:
    print(f"Error: input file must have atleast {sampleSize} lines, current file has {len(lines)}")
    exit(1)
  
  sampledLines = random.sample(lines, sampleSize)
  return sampledLines

def CreateFieldMap(title, sampledLines):
  fieldMap = {}
  for index, line in enumerate(sampledLines):
    fieldName = "FIELD_" + str(index + 1)
    fieldMap[fieldName] = line

  fieldMap["TITLE"] = title

  return fieldMap

def ReplaceInString(string, replacementMap):
  regexPattern = re.compile(r'\[\[(.*?)\]\]')

  def replacer(match):
    key = match.group(1)
    return replacementMap.get(key, match.group(0))
  
  return regexPattern.sub(replacer, string)


def ReplaceTableFields(doc, replacementMap):
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for para in cell.paragraphs:
            newText = ReplaceInString(para.text, replacementMap)
            para.clear()
            run = para.add_run(newText)
            run.font.name = "Comic Sans MS"
  return doc

def HandleArgs():
  parser = argparse.ArgumentParser(description= "Let's play bingo.")
  parser.add_argument("-i", required = True,  help = "A text file where each bingo option is on a separate line.")
  parser.add_argument("-o", required = True,  help ="The title of the bingo cards")
  parser.add_argument("-n", default = 1, type = int, required = False, help = "The amount of bingo cards to generate.")

  args = parser.parse_args()
  return args

if __name__ == "__main__":
  args = HandleArgs()
  nrOfCards = int(args.n)

  for i in range(nrOfCards):
      doc = Document("Bingo Template.docx")
      sampledLines = SampleInputTxt(args.i)
      fieldMap = CreateFieldMap(args.o, sampledLines)
      newDoc = ReplaceTableFields(doc, fieldMap)
      outputFileName = str(args.o) + "_BINGO_" + str(i + 1) + ".docx"
      newDoc.save(outputFileName)
